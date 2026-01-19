import os
import json
import io
from datetime import datetime
from typing import List, Optional

from fastapi import FastAPI, File, UploadFile, Request, HTTPException, status
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles # Although we use CDN, good practice
from google.cloud import firestore
import google.generativeai as genai
from pypdf import PdfReader # For PDF text extraction
from docx import Document as DocxDocument # For DOCX text extraction

# --- Environment Configuration ---
# Ensure GEMINI_API_KEY is set in your Cloud Run environment variables
# For local development, you can use python-dotenv:
# from dotenv import load_dotenv
# load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable not set.")

genai.configure(api_key=GEMINI_API_KEY)
GEMINI_MODEL_FLASH = 'gemini-1.5-flash' # Prioritizing Flash for cost/speed
LLM_MODEL = genai.GenerativeModel(GEMINI_MODEL_FLASH)

# Initialize Firestore DB client
# In Cloud Run, this will automatically pick up credentials from the service account
db = firestore.Client()
DOCUMENTS_COLLECTION = 'documents'

app = FastAPI(
    title="EntaGen: Анализ на Документи",
    description="Висококачествен инструмент за анализ на корпоративни документи с Gemini 1.5 Flash и FastAPI.",
    version="1.0.0"
)

# --- Utility Functions for Text Extraction ---

async def extract_text_from_pdf(file: UploadFile) -> str:
    """Извлича текст от PDF файл."""
    try:
        reader = PdfReader(io.BytesIO(await file.read()))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Неуспешно извличане на текст от PDF: {e}"
        )

async def extract_text_from_docx(file: UploadFile) -> str:
    """Извлича текст от DOCX файл."""
    try:
        document = DocxDocument(io.BytesIO(await file.read()))
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Неуспешно извличане на текст от DOCX: {e}"
        )

async def analyze_document_with_gemini(text_content: str) -> str:
    """Използва Gemini 1.5 Flash за генериране на резюме."""
    try:
        prompt = (
            "Вие сте експерт по корпоративен анализ на документи. "
            "Моля, резюмирайте следния документ на български език. "
            "Обобщението трябва да е кратко, стегнато и да улавя основните точки. "
            "Използвайте максимум 200 думи и форматирайте като списък с точки, ако е приложимо:\n\n"
            f"{text_content}"
        )
        response = await LLM_MODEL.generate_content_async(prompt)
        return response.text
    except Exception as e:
        # Log the error for debugging
        print(f"Грешка при анализ с Gemini: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Възникна грешка при анализа на документа. Моля, опитайте отново."
        )

# --- FastAPI Endpoints ---

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """
    Основна страница на приложението EntaGen.
    Предоставя HTML интерфейс за качване и преглед на документи.
    """
    return HTMLResponse(content=get_html_content(), status_code=status.HTTP_200_OK)

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Качва документ за анализ.
    Поддържа PDF и DOCX файлове.
    """
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не е избран файл."
        )

    file_extension = file.filename.split('.')[-1].lower()
    text_content = ""

    if file_extension == "pdf":
        text_content = await extract_text_from_pdf(file)
    elif file_extension == "docx":
        text_content = await extract_text_from_docx(file)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Поддържат се само PDF и DOCX файлове."
        )

    if not text_content.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Файлът е празен или не може да бъде извлечен текст."
        )

    summary = await analyze_document_with_gemini(text_content)

    try:
        doc_ref = db.collection(DOCUMENTS_COLLECTION).document()
        doc_data = {
            "name": file.filename,
            "summary": summary,
            "status": "Анализиран",
            "timestamp": firestore.SERVER_TIMESTAMP # Use server timestamp for consistency
        }
        await doc_ref.set(doc_data) # Await the async set operation

        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={"message": "Документът е успешно качен и анализиран.", "document_id": doc_ref.id}
        )
    except Exception as e:
        print(f"Грешка при запис във Firestore: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Грешка при запазване на документа в базата данни."
        )

@app.get("/documents", response_class=JSONResponse)
async def get_documents():
    """
    Връща списък с всички анализирани документи от Firestore.
    """
    try:
        documents = []
        docs = db.collection(DOCUMENTS_COLLECTION).order_by("timestamp", direction=firestore.Query.DESCENDING).stream()
        for doc in docs:
            doc_data = doc.to_dict()
            # Convert timestamp to a readable string if it exists
            if 'timestamp' in doc_data and doc_data['timestamp']:
                # firestore.SERVER_TIMESTAMP returns a Timestamp object
                doc_data['timestamp'] = doc_data['timestamp'].isoformat() if hasattr(doc_data['timestamp'], 'isoformat') else str(doc_data['timestamp'])
            
            documents.append({"id": doc.id, **doc_data})
        return JSONResponse(content=documents, status_code=status.HTTP_200_OK)
    except Exception as e:
        print(f"Грешка при извличане на документи от Firestore: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Грешка при извличане на документи."
        )

@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """
    Изтрива документ от Firestore по неговия ID.
    """
    try:
        doc_ref = db.collection(DOCUMENTS_COLLECTION).document(doc_id)
        if not (await doc_ref.get()).exists: # Check if document exists
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документът не е намерен."
            )
        await doc_ref.delete() # Await the async delete operation
        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={"message": "Документът е успешно изтрит."}
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        print(f"Грешка при изтриване на документ във Firestore: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Грешка при изтриване на документа."
        )

# --- HTML Frontend Content ---

def get_html_content():
    """
    Генерира HTML съдържанието за потребителския интерфейс.
    Използва Tailwind CSS (CDN) за модерен Dark Mode дизайн.
    """
    return f"""
<!DOCTYPE html>
<html lang="bg" class="dark">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>EntaGen: Анализ на Документи</title>
    <!-- Tailwind CSS CDN -->
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        /* Dark Mode configuration for Tailwind CSS */
        :root {{
            --color-primary: #1a73e8; /* Google Blue */
            --color-danger: #ea4335; /* Google Red */
        }}
        .dark {{
            color-scheme: dark;
        }}
        .dark body {{
            background-color: #1a202c; /* Dark Gray */
            color: #e2e8f0; /* Light Gray */
        }}
        .dark .bg-white {{
            background-color: #2d3748; /* Darker Gray */
        }}
        .dark .text-gray-700 {{
            color: #cbd5e0; /* Lighter Gray */
        }}
        .dark .border-gray-300 {{
            border-color: #4a5568; /* Gray Border */
        }}
        .dark .hover\\:bg-gray-50 {{
            background-color: #4a5568; /* Darker Gray on hover */
        }}
        /* Custom drag-and-drop styles */
        .drag-area {{
            border: 2px dashed #4a5568;
            background-color: #2d3748;
            transition: background-color 0.3s ease;
        }}
        .drag-area.highlight {{
            background-color: #4a5568;
            border-color: var(--color-primary);
        }}
        .upload-button {{
            background-color: var(--color-primary);
        }}
        .upload-button:hover {{
            background-color: #1558b3;
        }}
        .delete-button {{
            background-color: var(--color-danger);
        }}
        .delete-button:hover {{
            background-color: #d12a1c;
        }}
        .spinner {{
            border: 4px solid rgba(255, 255, 255, 0.3);
            border-top: 4px solid var(--color-primary);
            border-radius: 50%;
            width: 30px;
            height: 30px;
            animation: spin 1s linear infinite;
        }}
        @keyframes spin {{
            0% {{ transform: rotate(0deg); }}
            100% {{ transform: rotate(360deg); }}
        }}
    </style>
</head>
<body class="bg-gray-900 text-gray-100 font-sans leading-normal tracking-normal p-4">

    <div class="container mx-auto p-4 bg-gray-800 shadow-xl rounded-lg mt-8 max-w-4xl">
        <h1 class="text-4xl font-bold text-center mb-8 text-white">EntaGen: Анализ на Документи</h1>

        <!-- File Upload Section -->
        <div class="mb-8 p-6 bg-gray-700 rounded-lg shadow-md">
            <h2 class="text-2xl font-semibold mb-4 text-white">Качване на Документ</h2>
            <div id="drag-area" class="drag-area flex flex-col items-center justify-center p-10 text-center rounded-lg cursor-pointer transition-all duration-300 ease-in-out hover:border-blue-500 hover:bg-gray-600">
                <input type="file" id="fileInput" accept=".pdf,.docx" multiple class="hidden">
                <svg class="w-16 h-16 text-gray-400 mb-3" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M7 16a4 4 0 01-.88-7.903A5 5 0 0115.9 6L16 6a5 5 0 011 9.9M15 13l-3-3m0 0l-3 3m3-3v12"></path></svg>
                <p class="text-lg text-gray-300">Плъзнете и пуснете файлове тук или <span class="text-blue-400 font-medium hover:underline">кликнете за избор</span></p>
                <p class="text-sm text-gray-400 mt-1">(Поддържат се .pdf и .docx файлове)</p>
                <div id="fileListPreview" class="mt-4 w-full text-left text-gray-300"></div>
            </div>
            <div id="uploadStatus" class="mt-4 text-center font-medium"></div>
            <div id="loadingSpinner" class="hidden mt-4 flex justify-center">
                <div class="spinner"></div>
            </div>
        </div>

        <!-- Document List Section -->
        <div class="p-6 bg-gray-700 rounded-lg shadow-md">
            <h2 class="text-2xl font-semibold mb-4 text-white">Анализирани Документи</h2>
            <div id="documentsList" class="space-y-4">
                <p id="noDocumentsMessage" class="text-gray-400 text-center">Няма анализирани документи.</p>
                <!-- Documents will be loaded here by JavaScript -->
            </div>
        </div>
    </div>

    <script>
        const dragArea = document.getElementById('drag-area');
        const fileInput = document.getElementById('fileInput');
        const fileListPreview = document.getElementById('fileListPreview');
        const uploadStatus = document.getElementById('uploadStatus');
        const loadingSpinner = document.getElementById('loadingSpinner');
        const documentsList = document.getElementById('documentsList');
        const noDocumentsMessage = document.getElementById('noDocumentsMessage');

        // --- Drag and Drop Handlers ---
        dragArea.addEventListener('click', () => fileInput.click());
        fileInput.addEventListener('change', (e) => handleFiles(e.target.files));

        dragArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            dragArea.classList.add('highlight');
        });
        dragArea.addEventListener('dragleave', () => {
            dragArea.classList.remove('highlight');
        });
        dragArea.addEventListener('drop', (e) => {
            e.preventDefault();
            dragArea.classList.remove('highlight');
            handleFiles(e.dataTransfer.files);
        });

        // --- File Handling and Upload ---
        async function handleFiles(files) {
            fileListPreview.innerHTML = ''; // Clear previous previews
            uploadStatus.textContent = '';
            loadingSpinner.classList.remove('hidden');

            if (files.length === 0) {
                loadingSpinner.classList.add('hidden');
                return;
            }

            for (const file of files) {
                const li = document.createElement('div');
                li.className = 'flex items-center justify-between p-2 bg-gray-600 rounded-md mb-2';
                li.innerHTML = `
                    <span class="text-gray-200">${file.name}</span>
                    <span class="text-gray-400 text-sm">Качване...</span>
                `;
                fileListPreview.appendChild(li);

                if (!['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'].includes(file.type)) {
                    li.children[1].textContent = 'Неподдържан формат';
                    li.children[1].className = 'text-red-400 text-sm';
                    continue; // Skip unsupported files
                }

                const formData = new FormData();
                formData.append('file', file);

                try {
                    const response = await fetch('/upload', {
                        method: 'POST',
                        body: formData,
                    });
                    const data = await response.json();

                    if (response.ok) {
                        li.children[1].textContent = 'Успешно анализиран!';
                        li.children[1].className = 'text-green-400 text-sm';
                        fetchDocuments(); // Refresh the list after successful upload
                    } else {
                        li.children[1].textContent = `Грешка: ${data.detail || 'Неизвестна грешка'}`;
                        li.children[1].className = 'text-red-400 text-sm';
                    }
                } catch (error) {
                    console.error('Upload error:', error);
                    li.children[1].textContent = 'Грешка в мрежата или сървъра.';
                    li.children[1].className = 'text-red-400 text-sm';
                }
            }
            loadingSpinner.classList.add('hidden');
        }

        // --- Document List Management ---
        async function fetchDocuments() {
            loadingSpinner.classList.remove('hidden');
            documentsList.innerHTML = ''; // Clear current list
            try {
                const response = await fetch('/documents');
                const documents = await response.json();

                if (documents.length === 0) {
                    noDocumentsMessage.classList.remove('hidden');
                } else {
                    noDocumentsMessage.classList.add('hidden');
                    documents.forEach(doc => {
                        const docElement = document.createElement('div');
                        docElement.id = `doc-${doc.id}`;
                        docElement.className = 'bg-gray-800 p-4 rounded-lg shadow-md border border-gray-600 flex flex-col md:flex-row md:items-start justify-between gap-4';
                        
                        const timestamp = doc.timestamp ? new Date(doc.timestamp).toLocaleString('bg-BG') : 'Неизвестна дата';

                        docElement.innerHTML = `
                            <div class="flex-grow">
                                <h3 class="text-xl font-semibold text-white mb-2">${doc.name}</h3>
                                <p class="text-gray-300 text-sm mb-1">Статус: <span class="text-green-400">${doc.status}</span></p>
                                <p class="text-gray-400 text-xs mb-3">Качено на: ${timestamp}</p>
                                <div class="bg-gray-700 p-3 rounded-md text-gray-200 text-sm leading-relaxed max-h-48 overflow-y-auto">
                                    <h4 class="font-medium text-white mb-1">Резюме:</h4>
                                    <p>${doc.summary || 'Няма налично резюме.'}</p>
                                </div>
                            </div>
                            <div class="flex-shrink-0 mt-4 md:mt-0">
                                <button onclick="deleteDocument('${doc.id}')" class="delete-button text-white px-4 py-2 rounded-md transition-colors duration-200 text-sm">
                                    Изтрий
                                </button>
                            </div>
                        `;
                        documentsList.appendChild(docElement);
                    });
                }
            } catch (error) {
                console.error('Error fetching documents:', error);
                documentsList.innerHTML = `<p class="text-red-400 text-center">Грешка при зареждане на документи.</p>`;
                noDocumentsMessage.classList.add('hidden'); // Hide if there's an error message
            } finally {
                loadingSpinner.classList.add('hidden');
            }
        }

        async function deleteDocument(docId) {
            if (!confirm('Сигурни ли сте, че искате да изтриете този документ?')) {
                return;
            }

            loadingSpinner.classList.remove('hidden');
            try {
                const response = await fetch(`/documents/${docId}`, {
                    method: 'DELETE',
                });
                const data = await response.json();

                if (response.ok) {
                    alert(data.message);
                    fetchDocuments(); // Refresh the list
                } else {
                    alert(`Грешка при изтриване: ${data.detail || 'Неизвестна грешка'}`);
                }
            } catch (error) {
                console.error('Delete error:', error);
                alert('Грешка в мрежата или сървъра при изтриване.');
            } finally {
                loadingSpinner.classList.add('hidden');
            }
        }

        // --- Initial Load ---
        document.addEventListener('DOMContentLoaded', fetchDocuments);
    </script>
</body>
</html>
    """
